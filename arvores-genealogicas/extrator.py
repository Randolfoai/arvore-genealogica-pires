# -*- coding: utf-8 -*-
"""
Extrator DOCX -> dados.json para as arvores genealogicas da Familia Pires.

ETAPA 1: extracao fiel + relatorios de qualidade. Nada e inventado,
normalizado ou "corrigido": todo registro carrega texto_original e
paragrafo_idx (indice do paragrafo NO DOCX SANITIZADO lido por este script,
que NAO coincide com os indices dos relatorios do DocxPrep).

Cada registro tambem carrega arquivo_origem + paragrafo_idx_local: o nome do
arquivo .docx de onde ele veio e o indice do paragrafo DENTRO DAQUELE
ARQUIVO. Isso existe porque a obra e um unico livro dividido em dois DOCX
fisicos; paragrafo_idx sozinho, no modo combinado, e um indice GLOBAL
continuo entre os arquivos (nao reinicia em zero na segunda parte) -- para
achar o paragrafo abrindo o Word de UM dos arquivos, use arquivo_origem +
paragrafo_idx_local, nao paragrafo_idx.

Uso (arquivo unico, legado -- reprocessa so uma parte isolada):
    python extrator.py [caminho_docx] [dir_saida]

Uso (combinado -- trata varios DOCX como um fluxo continuo de paragrafos,
necessario quando um tronco-raiz comeca num arquivo e continua no outro):
    python extrator.py --combinado <docx1> <docx2> [... <docxN>] [dir_saida]

Sem argumentos, usa originais/primeira_parte_limpo.docx e o diretorio atual.
"""
from __future__ import annotations

import json
import re
import sys
import unicodedata
from dataclasses import dataclass, field, asdict
from pathlib import Path

# --------------------------------------------------------------------------
# Helpers de texto
# --------------------------------------------------------------------------

def desacentuar(s: str) -> str:
    return "".join(
        c for c in unicodedata.normalize("NFKD", s or "")
        if not unicodedata.combining(c)
    )


def norm(s: str) -> str:
    """minusculo, sem acento, espacos colapsados."""
    return re.sub(r"\s+", " ", desacentuar(s or "").lower()).strip()


# Palavras-numero usadas nas frases-resumo (contagens pequenas por extenso).
PALAVRA_NUM = {
    "um": 1, "uma": 1, "dois": 2, "duas": 2, "tres": 3, "quatro": 4,
    "cinco": 5, "seis": 6, "sete": 7, "oito": 8, "nove": 9, "dez": 10,
    "onze": 11, "doze": 12, "treze": 13, "quatorze": 14, "catorze": 14,
    "quinze": 15, "dezesseis": 16, "dezessete": 17, "dezoito": 18,
    "dezenove": 19, "vinte": 20,
}


def palavra_ou_num(token: str):
    """Converte '41' -> 41 ou 'sete' -> 7. Retorna (valor|None, texto_original)."""
    if token is None:
        return None
    t = token.strip()
    tnum = t.replace(".", "")
    if tnum.isdigit():
        return int(tnum)
    return PALAVRA_NUM.get(norm(t))


# Palavras que, num parenteses ou apos barra, indicam OBSERVACAO (nao apelido/conjuge).
OBS_KEYWORDS = (
    "faleceu", "solteir", "casamento", "gemeo", "natural", "viuv",
    "adotad", "crianca", "bebe", "sem informac", "nao ", "in memorian",
    "in memoriam", "falecid",
)

# Conectores de cabecalho de geracao (o "de/do/da" apos a palavra-nivel).
PREPOS = r"(?:de|do|da|dos|das)"

# Palavras-nivel (com typos conhecidos) -> nivel normalizado.
NIVEL_MAP = [
    (r"tataranetos?", "tataranetos"),
    (r"tetranetos?", "tetranetos"),
    (r"trinetos?", "trinetos"),
    (r"bisnetos?", "bisnetos"),
    (r"netos?", "netos"),
    (r"flhos?", "filhos"),
    (r"filhos?", "filhos"),
]
_CABEC_RE = re.compile(
    r"^\s*(?P<kw>tataranetos?|tetranetos?|trinetos?|bisnetos?|netos?|flhos?|filhos?)"
    r"\s+(?P<prep>de|do|da|dos|das)\b",
    re.IGNORECASE,
)


# --------------------------------------------------------------------------
# Detectores (funcoes puras, testaveis)
# --------------------------------------------------------------------------

def eh_marcador_foto(estilo: str, texto: str) -> bool:
    if (estilo or "") == "MARCADOR_FOTO":
        return True
    return bool(re.match(r"\s*\[?\s*fotos?\b", norm(texto))) and "[" in (texto or "")


def detectar_separador(texto: str):
    """Retorna 'tronco-raiz' | 'tronco' | 'patriarca' | None."""
    t = norm(texto)
    if t in ("tronco-raiz", "tronco raiz", "troncoraiz"):
        return "tronco-raiz"
    if t == "tronco":
        return "tronco"
    if t == "patriarca":
        return "patriarca"
    return None


def detectar_cabecalho(texto: str):
    """
    Detecta cabecalho de geracao por PALAVRA-CHAVE + FORMA (curto), qualquer estilo.
    Retorna dict {nivel, referencia, denominados, rotulo} ou None.
    Rejeita narrativa longa que apenas comeca com 'Filho de ...'.
    """
    if not texto:
        return None
    bruto = texto.strip()
    m = _CABEC_RE.match(bruto)
    if not m:
        return None

    baixa = norm(bruto)
    n_palavras = len(baixa.split())
    termina_dois_pontos = bruto.rstrip().endswith(":")
    tem_denominado = "denominado" in baixa

    # Guarda anti-narrativa: cabecalhos sao curtos. Aceita se termina em ':',
    # contem '(denominado...)', ou tem poucas palavras.
    if not (termina_dois_pontos or tem_denominado or n_palavras <= 12):
        return None

    kw = m.group("kw")
    nivel = None
    for pat, nome in NIVEL_MAP:
        if re.fullmatch(pat, kw, re.IGNORECASE):
            nivel = nome
            break

    # denominados: 'Tronco', 'Troncos', 'Tronco-raiz'
    denominados = None
    md = re.search(r"denominado[s]?\s+([\w\- ]+?)\s*\)", bruto, re.IGNORECASE)
    if md:
        denominados = md.group(1).strip()

    # referencia = miolo apos "<nivel> <prep>" ate ':' ou '(' ou frase de cauda.
    resto = bruto[m.end():].strip()
    resto = re.split(r"\s*\(", resto)[0]           # tira "(Denominados...)"
    resto = resto.split(":")[0]                     # tira ':'
    resto = re.sub(
        r"\s+e\s+(seus?|os)\s+respe[c]?tivos?\s+c[oô]njuges?.*$", "",
        resto, flags=re.IGNORECASE,
    )
    referencia = resto.strip(" ,.;")

    return {
        "nivel": nivel,
        "referencia": referencia,
        "denominados": denominados,
        "rotulo": bruto,
    }


def _classificar_parentese(conteudo: str) -> str:
    """Retorna 'apelido' ou 'observacao' para o texto dentro de (...)."""
    c = conteudo.strip()
    baixa = norm(c)
    if any(k in baixa for k in OBS_KEYWORDS):
        return "observacao"
    if any(ch.isdigit() for ch in c):
        return "observacao"
    # apelido: curto (ate 2 palavras) sem palavras de status
    if len(c.split()) <= 2:
        return "apelido"
    return "observacao"


def _limpar_terminador(s: str) -> str:
    """Remove terminadores de lista em prosa: '.', ',', ' e' no fim."""
    t = s.rstrip()
    mudou = True
    while mudou:
        mudou = False
        for suf in (".", ",", ";"):
            if t.endswith(suf):
                t = t[:-1].rstrip()
                mudou = True
        if re.search(r"\se$", t):
            t = re.sub(r"\se$", "", t).rstrip()
            mudou = True
    return t


def parse_pessoa(texto: str) -> dict:
    """
    Faz parse de uma linha de pessoa. Campos: nome, apelido, conjuge,
    observacao, confianca ('alta'|'media'). Nunca descarta: o texto sempre
    vira algo. texto_original e paragrafo_idx sao adicionados pelo chamador.
    """
    original = texto
    trabalho = _limpar_terminador(texto.strip())
    apelido = None
    conjuge = None
    observacoes = []
    incerto = False

    # separa nome / conjuge pela primeira barra
    if "/" in trabalho:
        esquerda, direita = trabalho.split("/", 1)
    else:
        esquerda, direita = trabalho, None

    # --- lado do nome: extrai parenteses ---
    def extrair_parenteses(seg: str):
        nonlocal incerto
        achados = re.findall(r"\(([^)]*)\)", seg)
        base = re.sub(r"\s*\([^)]*\)", "", seg).strip()
        aps, obs = [], []
        for a in achados:
            tipo = _classificar_parentese(a)
            if tipo == "apelido":
                aps.append(a.strip())
            else:
                obs.append(a.strip())
        # parentese nao fechado -> incerteza
        if seg.count("(") != seg.count(")"):
            incerto = True
        return base, aps, obs

    nome, aps_nome, obs_nome = extrair_parenteses(esquerda)
    if aps_nome:
        apelido = aps_nome[0]
        if len(aps_nome) > 1:
            observacoes.extend(aps_nome[1:])
            incerto = True
    observacoes.extend(obs_nome)

    # --- lado do conjuge ---
    if direita is not None:
        dbase, aps_c, obs_c = extrair_parenteses(direita)
        dbase = dbase.strip()
        observacoes.extend(obs_c)
        observacoes.extend(aps_c)  # parentese do conjuge -> observacao
        if dbase == "":
            # barra sem nome (ex.: "/ faleceu ainda crianca," ja foi p/ obs)
            conjuge = None
            if not obs_c:
                incerto = True
        elif any(k in norm(dbase) for k in OBS_KEYWORDS):
            observacoes.append(dbase)     # ex.: "Solteira", "faleceu..."
            conjuge = None
        else:
            conjuge = dbase

    nome = nome.strip(" ,.;")
    observacao = "; ".join(o for o in observacoes if o) or None

    # confianca
    confianca = "alta"
    if observacao or incerto:
        confianca = "media"
    if not nome or len(nome) < 2:
        confianca = "media"
    # nome com muitas palavras suspeito (pode ser narrativa/caption)
    if len(nome.split()) > 8:
        confianca = "media"

    return {
        "nome": nome,
        "apelido": apelido,
        "conjuge": conjuge,
        "observacao": observacao,
        "confianca": confianca,
        "texto_original": original,
    }


_RESUMO_RE_TOTAL = re.compile(r"(\d[\d\.]*)\s+membros\s+descendentes", re.IGNORECASE)


def parse_resumo(texto: str):
    """
    Extrai numeros declarados pelo autor da frase-resumo de um tronco-raiz.
    Retorna dict ou None se nao for uma frase-resumo.
    """
    baixa = norm(texto)
    if "membros descendentes" not in baixa:
        return None

    def achar(padrao):
        m = re.search(padrao, texto, re.IGNORECASE)
        return palavra_ou_num(m.group(1)) if m else None

    total = None
    mt = _RESUMO_RE_TOTAL.search(texto)
    if mt:
        total = palavra_ou_num(mt.group(1))

    # casal: entre "desse [T/t]ronco[- ]raiz," e ", sendo"
    casal = None
    mc = re.search(
        r"desse\s+tronco[- ]?raiz,?\s*(.+?)\s*,?\s+sendo\b",
        texto, re.IGNORECASE,
    )
    if mc:
        casal = mc.group(1).strip()

    filhos = achar(r"sendo\s+([\wçãáéêíóôúà]+|\d+)\s+filhos")
    netos = achar(r"([\wçãáéêíóôúà]+|\d+)\s+netos")
    bisnetos = achar(r"([\wçãáéêíóôúà]+|\d+)\s+bisnetos")
    trinetos = achar(r"([\wçãáéêíóôúà]+|\d+)\s+trinetos")
    tataranetos = achar(r"([\wçãáéêíóôúà]+|\d+)\s+tataranetos")

    return {
        "descendentes_total": total,
        "casal": casal,
        "filhos": filhos,
        "netos": netos,
        "bisnetos": bisnetos,
        "trinetos": trinetos,
        "tataranetos": tataranetos,
        "trecho_original": texto.strip(),
    }


def parece_pessoa(texto: str) -> bool:
    """Heuristica leve: linha de pessoa e curta e nao e legenda/narrativa."""
    t = texto.strip()
    baixa = norm(t)
    if not t:
        return False
    if baixa.startswith("legenda"):
        return False
    if len(t.split()) > 20:
        return False
    return True


def eh_nota_nao_pessoa(texto: str) -> bool:
    """Notas editoriais/observacoes que nao sao pessoas."""
    low = norm(texto)
    if low.startswith("orelha"):
        return True
    if low.startswith("nao tiveram") or low.startswith("nao teve"):
        return True
    if low in ("sem filhos", "sem descendentes"):
        return True
    if low.startswith("sem informac"):
        return True
    return False


def eh_legenda_apos_foto(estilo: str, texto: str) -> bool:
    """
    Detecta legenda/caption que aparece logo apos um [FOTO], em QUALQUER estilo.
    - sequencia de primeiros-nomes soltos (sem 'de/da/do', sem virgula/barra);
    - ou frase descritiva (com/na/em/ano/numeros).
    So deve ser chamada quando o paragrafo anterior foi um marcador de foto.
    """
    s = texto.strip()
    low = norm(s)
    if "/" in s:
        return False
    toks = low.split()
    tem_conector = any(t in ("de", "da", "do", "dos", "das") for t in toks)
    if ("," not in s) and len(toks) >= 3 and not tem_conector and estilo != "Body Text":
        return True
    if re.search(r"\b(com|na|no|em|decada|janeiro|fevereiro|marco|abril|maio|"
                 r"junho|julho|agosto|setembro|outubro|novembro|dezembro|ano|"
                 r"foto|legenda|montado|fazenda)\b", low):
        return True
    if any(ch.isdigit() for ch in s):
        return True
    return False


def eh_legenda_sem_conector(estilo: str, texto: str) -> bool:
    """
    Complementa eh_legenda_apos_foto: pega legenda solta pos-[FOTO] no
    formato "Nome1 Nome2 Nome3 ..." (varios primeiros-nomes soltos, sem
    virgula, sem barra, sem conector 'de/da/do', sem palavra-gatilho de
    data/local/foto/legenda, e por isso NAO capturada por
    eh_legenda_apos_foto -- cujo primeiro ramo so dispara com
    estilo != "Body Text"). Achado em auditoria (2026-08-15): 7 casos assim
    viravam pessoa falsa (nome de varias palavras sem nenhuma delas ser
    conector). So deve ser chamada quando o paragrafo anterior foi [FOTO] e
    eh_legenda_apos_foto ja retornou False.
    """
    s = texto.strip()
    if "/" in s or "," in s:
        return False
    low = norm(s)
    toks = low.split()
    if len(toks) < 3:
        return False
    tem_conector = any(t in ("de", "da", "do", "dos", "das") for t in toks)
    if tem_conector:
        return False
    palavras = s.split()
    capitalizadas = sum(1 for w in palavras if w[:1].isupper())
    return capitalizadas >= max(2, len(palavras) - 1)


def multiplas_pessoas_no_nome(nome: str) -> bool:
    """Um nome ja parseado nao deveria conter virgula; se contem, ha 2+ pessoas."""
    return "," in (nome or "")


# --------------------------------------------------------------------------
# Modelo de dados
# --------------------------------------------------------------------------

@dataclass
class Pessoa:
    nome: str
    apelido: str | None
    conjuge: str | None
    observacao: str | None
    confianca: str
    texto_original: str
    paragrafo_idx: int
    arquivo_origem: str
    paragrafo_idx_local: int


@dataclass
class Geracao:
    nivel: str | None
    rotulo_original: str
    referencia: str | None
    paragrafo_idx: int
    arquivo_origem: str
    paragrafo_idx_local: int
    pessoas: list = field(default_factory=list)


@dataclass
class TroncoFamiliar:
    id: str
    rotulo_separador: str
    paragrafo_idx: int
    arquivo_origem: str
    paragrafo_idx_local: int
    casal: list = field(default_factory=list)      # [{texto_original, paragrafo_idx, arquivo_origem, paragrafo_idx_local}]
    biografia_idx: list = field(default_factory=list)
    geracoes: list = field(default_factory=list)


@dataclass
class TroncoRaiz:
    id: str
    rotulo_separador: str
    paragrafo_idx: int
    arquivo_origem: str
    paragrafo_idx_local: int
    casal: list = field(default_factory=list)
    resumo_autor: dict | None = None
    biografia_idx: list = field(default_factory=list)
    primeira_geracao: dict | None = None            # {rotulo, idx, pessoas[]}
    troncos_familiares: list = field(default_factory=list)


@dataclass
class Patriarca:
    id: str
    rotulo_separador: str
    paragrafo_idx: int
    arquivo_origem: str
    paragrafo_idx_local: int
    identificacao: list = field(default_factory=list)
    geracoes: list = field(default_factory=list)


# --------------------------------------------------------------------------
# Extrator (maquina de estados)
# --------------------------------------------------------------------------

class Extrator:
    def __init__(self, arquivo_origem: str, origem_por_idx: dict | None = None,
                 arquivos_origem: list | None = None):
        self.arquivo_origem = arquivo_origem
        # origem_por_idx: {indice_global: (arquivo, indice_local)}. Se vazio,
        # assume-se arquivo unico e indice local == indice global.
        self.origem_por_idx = origem_por_idx or {}
        self.arquivos_origem = arquivos_origem or [arquivo_origem]
        self.patriarcas: list[Patriarca] = []
        self.troncos_raiz: list[TroncoRaiz] = []
        self.ambiguidades: list[dict] = []
        self.fotos: list[dict] = []

        # estado corrente
        self.pat: Patriarca | None = None
        self.raiz: TroncoRaiz | None = None
        self.tf: TroncoFamiliar | None = None
        self.geracao: Geracao | None = None
        self.escopo: str = "nenhum"   # 'primeira_geracao' | 'geracao_tf' | 'geracao_pat'
        self.esperando_casal: str | None = None   # 'raiz' | 'tf' | None
        self.ultimo_foi_foto = False

        self._c_raiz = 0
        self._c_tf = 0
        self._c_pat = 0

    # --- origem (arquivo + indice local) de um indice global ---
    def _origem(self, i):
        if i in self.origem_por_idx:
            return self.origem_por_idx[i]
        return (self.arquivo_origem, i)

    # --- registro de ambiguidades ---
    def _amb(self, tipo, idx, texto, nota=""):
        arq, local = self._origem(idx)
        self.ambiguidades.append({
            "tipo": tipo, "paragrafo_idx": idx,
            "arquivo_origem": arq, "paragrafo_idx_local": local,
            "texto_original": texto, "nota": nota,
        })

    def processar(self, paras):
        for i, p in enumerate(paras):
            estilo = p.style.name if p.style else ""
            texto = p.text or ""
            if texto.strip() == "":
                continue

            if eh_marcador_foto(estilo, texto):
                arq, local = self._origem(i)
                self.fotos.append({"paragrafo_idx": i, "arquivo_origem": arq,
                                   "paragrafo_idx_local": local, "marcador": texto.strip()})
                self.ultimo_foi_foto = True
                continue

            sep = detectar_separador(texto)
            if sep:
                self._sep(sep, i, texto)
                self.ultimo_foi_foto = False
                continue

            cab = detectar_cabecalho(texto)
            if cab:
                self._cabecalho(cab, i, texto)
                self.ultimo_foi_foto = False
                continue

            # linha de casal logo apos separador?
            if self.esperando_casal and self._parece_casal(estilo, texto):
                self._casal(i, texto)
                self.ultimo_foi_foto = False
                continue
            else:
                self.esperando_casal = None

            # dentro de uma geracao: candidato a pessoa
            if self.escopo in ("primeira_geracao", "geracao_tf", "geracao_pat"):
                if eh_nota_nao_pessoa(texto):
                    self._amb("nota_nao_pessoa", i, texto,
                              "Nota editorial/observacao; nao e pessoa.")
                    self.ultimo_foi_foto = False
                    continue
                # legenda logo apos foto (qualquer estilo) -> caption, nao pessoa
                if self.ultimo_foi_foto and eh_legenda_apos_foto(estilo, texto):
                    self._amb("legenda_apos_foto", i, texto,
                              "Logo apos [FOTO]; tratado como legenda/caption.")
                    self.ultimo_foi_foto = False
                    continue
                # legenda solta pos-foto sem conector 'de/da/do' (ponto cego
                # de eh_legenda_apos_foto quando o estilo e "Body Text" e nao
                # ha palavra-gatilho) -> tambem caption, nao pessoa
                if self.ultimo_foi_foto and eh_legenda_sem_conector(estilo, texto):
                    self._amb("legenda_sem_conector", i, texto,
                              "Logo apos [FOTO]; lista solta de nomes sem "
                              "conector 'de/da/do', virgula ou barra -- nao "
                              "capturada por eh_legenda_apos_foto.")
                    self.ultimo_foi_foto = False
                    continue
                if not parece_pessoa(texto):
                    self._amb("item_vs_narrativa", i, texto,
                              "Dentro de geracao mas com forma de narrativa/legenda; nao virou pessoa.")
                    self.ultimo_foi_foto = False
                    continue
                d = parse_pessoa(texto)
                d["paragrafo_idx"] = i
                d["arquivo_origem"], d["paragrafo_idx_local"] = self._origem(i)
                if not d["nome"]:
                    self._amb("pessoa_sem_nome", i, texto,
                              "Parse nao encontrou nome (linha so com parenteses/pontuacao).")
                    self.ultimo_foi_foto = False
                    continue
                if multiplas_pessoas_no_nome(d["nome"]):
                    self._amb("possivel_multiplas_pessoas", i, texto,
                              "Possivel mais de uma pessoa no mesmo paragrafo (virgula interna).")
                    d["confianca"] = "media"
                self._pessoa_dict(d)
                self.ultimo_foi_foto = False
                continue

            # fora de geracao: narrativa/biografia -> guarda o indice
            self._biografia(i, texto)
            self.ultimo_foi_foto = False

    # --- handlers ---
    def _sep(self, sep, i, texto):
        arq, local = self._origem(i)
        if sep == "patriarca":
            self._c_pat += 1
            self.pat = Patriarca(f"patriarca-{self._c_pat:02d}", texto.strip(), i, arq, local)
            self.patriarcas.append(self.pat)
            self.raiz = self.tf = self.geracao = None
            self.escopo = "nenhum"
            self.esperando_casal = None
        elif sep == "tronco-raiz":
            self._c_raiz += 1
            self.raiz = TroncoRaiz(f"tronco-raiz-{self._c_raiz:02d}", texto.strip(), i, arq, local)
            self.troncos_raiz.append(self.raiz)
            self.tf = self.geracao = None
            self.escopo = "nenhum"
            self.esperando_casal = "raiz"
        elif sep == "tronco":
            if self.raiz is None:
                self._amb("tronco_sem_raiz", i, texto,
                          "Separador 'Tronco' sem Tronco-raiz corrente.")
                return
            self._c_tf += 1
            self.tf = TroncoFamiliar(
                f"{self.raiz.id}-tf-{self._c_tf:02d}", texto.strip(), i, arq, local)
            self.raiz.troncos_familiares.append(self.tf)
            self.geracao = None
            self.escopo = "nenhum"
            self.esperando_casal = "tf"

    def _cabecalho(self, cab, i, texto):
        arq, local = self._origem(i)
        ger = Geracao(cab["nivel"], cab["rotulo"], cab["referencia"], i, arq, local)
        self.esperando_casal = None
        if self.tf is not None:
            self.tf.geracoes.append(ger)
            self.geracao = ger
            self.escopo = "geracao_tf"
        elif self.raiz is not None:
            # primeira geracao da raiz (os filhos = troncos)
            if self.raiz.primeira_geracao is None:
                self.raiz.primeira_geracao = {
                    "nivel": cab["nivel"], "rotulo_original": cab["rotulo"],
                    "referencia": cab["referencia"], "denominados": cab["denominados"],
                    "paragrafo_idx": i, "arquivo_origem": arq, "paragrafo_idx_local": local,
                    "pessoas": [],
                }
                self._pg = self.raiz.primeira_geracao
                self.escopo = "primeira_geracao"
                self.geracao = None
            else:
                # cabecalho extra na raiz antes de qualquer Tronco -> reporta e anexa
                self._amb("cabecalho_extra_na_raiz", i, texto,
                          "Segundo cabecalho na raiz antes do primeiro 'Tronco'.")
                self.tf = None
                self.geracao = ger
                # cria um pseudo-container na raiz
                self.raiz.__dict__.setdefault("geracoes_extra", []).append(ger)
                self.escopo = "geracao_tf"
                self._geracao_alvo = ger
        elif self.pat is not None:
            self.pat.geracoes.append(ger)
            self.geracao = ger
            self.escopo = "geracao_pat"
        else:
            self._amb("cabecalho_sem_contexto", i, texto,
                      "Cabecalho de geracao sem patriarca/raiz/tronco corrente.")

    def _parece_casal(self, estilo, texto):
        baixa = norm(texto)
        if estilo in ("Normal", "Título 31", "Titulo 31"):
            return True
        if baixa.startswith("esposo") or baixa.startswith("esposa"):
            return True
        if "(filh" in baixa or "(filha" in baixa or "(filho" in baixa:
            return True
        return False

    def _casal(self, i, texto):
        arq, local = self._origem(i)
        item = {"texto_original": texto.strip(), "paragrafo_idx": i,
                "arquivo_origem": arq, "paragrafo_idx_local": local}
        if self.esperando_casal == "raiz" and self.raiz is not None:
            self.raiz.casal.append(item)
        elif self.esperando_casal == "tf" and self.tf is not None:
            self.tf.casal.append(item)

    def _pessoa_dict(self, d):
        pessoa = Pessoa(**{k: d[k] for k in (
            "nome", "apelido", "conjuge", "observacao",
            "confianca", "texto_original", "paragrafo_idx",
            "arquivo_origem", "paragrafo_idx_local")})
        if self.escopo == "primeira_geracao":
            self._pg["pessoas"].append(asdict(pessoa))
        else:
            self.geracao.pessoas.append(asdict(pessoa))

    def _biografia(self, i, texto):
        alvo = self.tf or self.raiz
        if alvo is not None:
            alvo.biografia_idx.append(i)
        # se for uma frase-resumo do tronco-raiz, captura
        r = parse_resumo(texto)
        if r and self.raiz is not None and self.raiz.resumo_autor is None:
            r["paragrafo_idx"] = i
            r["arquivo_origem"], r["paragrafo_idx_local"] = self._origem(i)
            self.raiz.resumo_autor = r

    # --- serializacao ---
    def to_dict(self):
        def raiz_dict(r: TroncoRaiz):
            d = asdict(r)
            return d
        return {
            "meta": {
                "arquivo_origem": self.arquivo_origem,
                "arquivos_origem": self.arquivos_origem,
                "nota_indices": (
                    "paragrafo_idx e indice GLOBAL, continuo entre os arquivos de "
                    "'arquivos_origem' na ordem listada (nao reinicia por arquivo). "
                    "Para achar o paragrafo no Word original, use arquivo_origem + "
                    "paragrafo_idx_local (indice DENTRO daquele arquivo), presentes "
                    "em cada registro."
                ),
            },
            "patriarcas": [asdict(p) for p in self.patriarcas],
            "troncos_raiz": [raiz_dict(r) for r in self.troncos_raiz],
            "fotos": self.fotos,
            "ambiguidades": self.ambiguidades,
        }


# --------------------------------------------------------------------------
# Cross-validacao para o relatorio de extracao
# --------------------------------------------------------------------------

def _nomes_set(*textos):
    s = set()
    for t in textos:
        for tok in re.split(r"[\s,]+", norm(t or "")):
            if len(tok) >= 3 and tok not in ("de", "do", "da", "dos", "das", "e"):
                s.add(tok)
    return s


NIVEL_ABS = {  # nivel relativo ao tronco familiar -> nivel absoluto na raiz
    "filhos": "netos",
    "netos": "bisnetos",
    "bisnetos": "trinetos",
    "trinetos": "tataranetos",
    "tataranetos": "tetranetos",
}


def contar_por_nivel(raiz: TroncoRaiz):
    """
    Contagem best-effort por nivel ABSOLUTO da raiz, para comparar com o resumo.
    - filhos = pessoas da primeira_geracao (ou nro de troncos familiares).
    - dentro de cada tronco familiar, cabecalhos 'guarda-chuva' (referencia ~ casal
      do tronco) marcam a mudanca de nivel; subgrupos herdam o nivel corrente.
    Retorna (contagens: dict nivel->int, metodo_notas: list[str]).
    """
    cont = {"filhos": 0, "netos": 0, "bisnetos": 0,
            "trinetos": 0, "tataranetos": 0, "tetranetos": 0}
    notas = []

    if raiz.primeira_geracao:
        cont["filhos"] += len(raiz.primeira_geracao["pessoas"])

    for tf in raiz.troncos_familiares:
        casal_nomes = _nomes_set(*[c["texto_original"] for c in tf.casal],
                                 tf.rotulo_separador)
        nivel_corrente = None
        for ger in tf.geracoes:
            ref_nomes = _nomes_set(ger.referencia or "")
            eh_guardachuva = bool(casal_nomes & ref_nomes) or (nivel_corrente is None)
            if eh_guardachuva and ger.nivel in NIVEL_ABS:
                nivel_corrente = NIVEL_ABS[ger.nivel]
            destino = nivel_corrente or NIVEL_ABS.get(ger.nivel or "", "netos")
            cont[destino] = cont.get(destino, 0) + len(ger.pessoas)
    return cont, notas


# --------------------------------------------------------------------------
# Relatorios
# --------------------------------------------------------------------------

def total_pessoas_raiz(raiz: TroncoRaiz):
    total = 0
    conf = {"alta": 0, "media": 0}
    if raiz.primeira_geracao:
        for p in raiz.primeira_geracao["pessoas"]:
            total += 1
            conf[p["confianca"]] = conf.get(p["confianca"], 0) + 1
    for tf in raiz.troncos_familiares:
        for ger in tf.geracoes:
            for p in ger.pessoas:
                total += 1
                conf[p["confianca"]] = conf.get(p["confianca"], 0) + 1
    return total, conf


def contar_registros(ext: Extrator):
    """
    Conta registros de pessoa, casal e cabecalho de geracao em todo o extrator
    (patriarcas + troncos-raiz + troncos familiares). Usado no relatorio para
    dar uma contagem total, objetiva, de rastreabilidade.
    """
    n_pessoa = n_casal = n_geracao = 0
    for p in ext.patriarcas:
        for g in p.geracoes:
            n_geracao += 1
            n_pessoa += len(g.pessoas)
    for r in ext.troncos_raiz:
        n_casal += len(r.casal)
        if r.primeira_geracao:
            n_geracao += 1
            n_pessoa += len(r.primeira_geracao["pessoas"])
        for tf in r.troncos_familiares:
            n_casal += len(tf.casal)
            for g in tf.geracoes:
                n_geracao += 1
                n_pessoa += len(g.pessoas)
    return {
        "pessoa": n_pessoa, "casal": n_casal, "geracao": n_geracao,
        "pessoa_mais_casal": n_pessoa + n_casal,
        "total_com_geracao": n_pessoa + n_casal + n_geracao,
    }


def relatorio_extracao(ext: Extrator) -> str:
    L = []
    L.append("# Relatório de extração — Árvores genealógicas (Família Pires)\n")
    L.append(f"Arquivo de origem: `{ext.arquivo_origem}`\n")
    L.append("> Índices de parágrafo referem-se ao DOCX sanitizado lido por este "
             "script (não coincidem com os índices dos relatórios do DocxPrep).\n")

    L.append("## ⚠️ Como ler as divergências (leia antes das tabelas)\n")
    L.append("1. **Os totais do autor são do tronco-raiz no LIVRO INTEIRO** (ambas as "
             "partes). Esta extração cobre **só a primeira parte**, que pode conter "
             "apenas um subconjunto dos descendentes de cada tronco-raiz — logo, "
             "um total extraído **menor** que o declarado é esperado, não é erro.")
    L.append("2. **O nível absoluto (netos/bisnetos/…) é uma estimativa.** No documento "
             "os cabeçalhos são relativos (\"Filhos de Conceição\", \"Netos de Conceição\") "
             "e aninhados por sub-grupos (\"Filhos de Leoiza\"). O mapeamento "
             "relativo→absoluto é heurístico (casa o \"de X\" com o casal do tronco) e "
             "**requer validação humana**. A contagem fiel por rótulo está no detalhamento "
             "por tronco familiar, mais abaixo.")
    L.append("3. A contagem de **filhos** e o **nº de troncos familiares** são as "
             "comparações mais confiáveis nesta parte.\n")

    L.append("## Cobertura de troncos-raiz\n")
    L.append(f"- Patriarcas encontrados: **{len(ext.patriarcas)}**")
    for p in ext.patriarcas:
        L.append(f"  - {p.id} (par. {p.paragrafo_idx})")
    L.append(f"- Troncos-raiz encontrados: **{len(ext.troncos_raiz)}**")
    L.append("- Referência (narrativa mestra do livro): 2 patriarcas, "
             "17 troncos-raízes e 70 troncos familiares no total da obra. "
             "A **primeira parte** contém apenas parte deles; a lista autoritativa "
             "\"Troncos familiares participantes do V Encontro\" não foi localizada "
             "nesta parte (provavelmente na segunda).\n")

    L.append("## Por tronco-raiz: extraído vs. declarado pelo autor\n")
    for r in ext.troncos_raiz:
        casal = "; ".join(c["texto_original"] for c in r.casal) or "—"
        L.append(f"### {r.id} — par. {r.paragrafo_idx}")
        L.append(f"- Casal-raiz (parágrafos capturados): {casal}")
        n_tf = len(r.troncos_familiares)
        n_pg = len(r.primeira_geracao["pessoas"]) if r.primeira_geracao else 0
        L.append(f"- Troncos familiares (sub-seções 'Tronco'): **{n_tf}**")
        L.append(f"- 1ª geração (filhos listados em "
                 f"'{(r.primeira_geracao or {}).get('rotulo_original','—')}'): **{n_pg}**")

        cont, _ = contar_por_nivel(r)
        total, conf = total_pessoas_raiz(r)

        if r.resumo_autor:
            ra = r.resumo_autor
            L.append(f"- Resumo do autor (par. {ra['paragrafo_idx']}): "
                     f"total={ra['descendentes_total']}, filhos={ra['filhos']}, "
                     f"netos={ra['netos']}, bisnetos={ra['bisnetos']}, "
                     f"trinetos={ra['trinetos']}, tataranetos={ra['tataranetos']}")
            L.append("")
            L.append("| Nível | Extraído (parte 1, estimativa) | Declarado (livro) |")
            L.append("|---|---|---|")
            pares = [
                ("filhos", cont["filhos"], ra["filhos"]),
                ("netos", cont["netos"], ra["netos"]),
                ("bisnetos", cont["bisnetos"], ra["bisnetos"]),
                ("trinetos", cont["trinetos"], ra["trinetos"]),
                ("tataranetos", cont["tataranetos"], ra["tataranetos"]),
            ]
            for nome, ex, dec in pares:
                L.append(f"| {nome} | {ex} | {dec if dec is not None else '—'} |")
            dec_total = ra["descendentes_total"]
            L.append(f"| **total** | **{total}** | **{dec_total if dec_total is not None else '—'}** |")
            if dec_total is not None and total == dec_total:
                L.append("\n✅ Total da parte 1 == total do livro: este tronco-raiz "
                         "parece **completo** na primeira parte.")
            else:
                L.append("\nℹ️ Total menor que o do livro: parte dos descendentes "
                         "deve estar na segunda parte (esperado).")
        else:
            L.append(f"- **Resumo do autor não encontrado** para este tronco-raiz.")
            L.append(f"- Total de pessoas extraídas: {total}")
        L.append(f"- Confiança: alta={conf['alta']}, média={conf['media']}")

        # detalhamento fiel por tronco familiar (por rótulo)
        L.append("")
        L.append("<details><summary>Detalhamento fiel por tronco familiar (contagem por rótulo)</summary>\n")
        if r.primeira_geracao:
            L.append(f"- **1ª geração** — `{r.primeira_geracao['rotulo_original']}` "
                     f"(par. {r.primeira_geracao['paragrafo_idx']}): "
                     f"{len(r.primeira_geracao['pessoas'])} pessoas")
        for tf in r.troncos_familiares:
            casal0 = tf.casal[0]["texto_original"] if tf.casal else "—"
            n_tf_pessoas = sum(len(g.pessoas) for g in tf.geracoes)
            L.append(f"- **{tf.id}** (sep. par. {tf.paragrafo_idx}) — {casal0}")
            L.append(f"  - gerações: {len(tf.geracoes)}, pessoas: {n_tf_pessoas}")
            for g in tf.geracoes:
                L.append(f"    - `{g.rotulo_original}` (par. {g.paragrafo_idx}): "
                         f"{len(g.pessoas)} pessoas")
        L.append("\n</details>")
        L.append("")

    # totais gerais
    tot_alta = tot_media = tot_geral = 0
    for r in ext.troncos_raiz:
        t, c = total_pessoas_raiz(r)
        tot_geral += t
        tot_alta += c["alta"]
        tot_media += c["media"]
    L.append("## Totais gerais\n")
    L.append(f"- Pessoas extraídas (todos os troncos-raiz): **{tot_geral}**")
    L.append(f"- Confiança alta: {tot_alta} · média: {tot_media}")
    L.append(f"- Marcadores de foto ignorados: {len(ext.fotos)}")
    L.append(f"- Casos em relatorio_ambiguidades.md: {len(ext.ambiguidades)}")
    L.append("")

    reg = contar_registros(ext)
    L.append("## Contagem de registros (pessoa/casal), para conferência de rastreabilidade\n")
    L.append("> Inclui patriarcas + todos os troncos-raiz. \"Registro\" aqui é qualquer "
             "item com `texto_original` e `paragrafo_idx` próprios — pessoa e linha de "
             "casal. Cabeçalhos de geração (\"Filhos de…\", \"Netos de…\") são contados "
             "à parte, pois usam `rotulo_original` em vez de `texto_original`.\n")
    L.append(f"- Registros de **pessoa**: {reg['pessoa']}")
    L.append(f"- Registros de **casal** (linhas do casal-raiz/casal-tronco): {reg['casal']}")
    L.append(f"- **Pessoa + casal**: **{reg['pessoa_mais_casal']}**")
    L.append(f"- Cabeçalhos de **geração** (contados à parte): {reg['geracao']}")
    L.append(f"- Pessoa + casal + geração (todos os registros com algum texto de origem "
             f"+ índice): **{reg['total_com_geracao']}**")
    L.append("")
    L.append("> **Como usar estes números:** confie em *filhos* e no *nº de troncos "
             "familiares* para conferir cobertura da parte 1. As colunas de "
             "netos/bisnetos/… são **estimativas** (ver aviso no topo) e servem de "
             "oráculo: onde a estimativa destoa muito do detalhamento por rótulo, há "
             "um cabeçalho relativo/aninhado a validar. O `relatorio_ambiguidades.md` "
             "lista cada parágrafo duvidoso com texto e índice.")
    L.append("")
    return "\n".join(L)


def relatorio_ambiguidades(ext: Extrator) -> str:
    L = []
    L.append("# Relatório de ambiguidades — Árvores genealógicas (Família Pires)\n")
    L.append(f"Arquivo de origem: `{ext.arquivo_origem}`\n")
    L.append(f"Total de casos: **{len(ext.ambiguidades)}**\n")

    por_tipo = {}
    for a in ext.ambiguidades:
        por_tipo.setdefault(a["tipo"], []).append(a)

    # pessoas de confianca media (revisar apelido vs observacao, conjuge, etc.)
    medias = []
    for r in ext.troncos_raiz:
        blocos = []
        if r.primeira_geracao:
            blocos.append(("primeira_geracao", r.primeira_geracao["pessoas"]))
        for tf in r.troncos_familiares:
            for ger in tf.geracoes:
                blocos.append((ger.rotulo_original, ger.pessoas))
        for rotulo, pessoas in blocos:
            for p in pessoas:
                if p["confianca"] == "media":
                    medias.append((r.id, rotulo, p))

    L.append("## Resumo por tipo\n")
    for tipo, itens in sorted(por_tipo.items()):
        L.append(f"- `{tipo}`: {len(itens)}")
    L.append(f"- `pessoa_confianca_media`: {len(medias)}")
    L.append("")

    for tipo, itens in sorted(por_tipo.items()):
        L.append(f"## {tipo} ({len(itens)})\n")
        for a in itens:
            L.append(f"- **par. {a['paragrafo_idx']}** "
                     f"(`{a['arquivo_origem']}` par. local {a['paragrafo_idx_local']}) "
                     f"— {a['nota']}")
            L.append(f"  - texto: `{a['texto_original'].strip()}`")
        L.append("")

    L.append(f"## pessoa_confianca_media ({len(medias)})\n")
    L.append("Parse aceito, mas com apelido/observação/cônjuge incertos — revisar.\n")
    for rid, rotulo, p in medias:
        L.append(f"- **par. {p['paragrafo_idx']}** "
                 f"(`{p['arquivo_origem']}` par. local {p['paragrafo_idx_local']}) "
                 f"({rid} / {rotulo})")
        L.append(f"  - nome=`{p['nome']}` · apelido=`{p['apelido']}` · "
                 f"cônjuge=`{p['conjuge']}` · obs=`{p['observacao']}`")
        L.append(f"  - texto: `{p['texto_original'].strip()}`")
    L.append("")
    return "\n".join(L)


# --------------------------------------------------------------------------
# main
# --------------------------------------------------------------------------

def _ler_paragrafos_com_origem(docx_paths: list[Path]):
    """
    Le varios .docx e concatena seus paragrafos num unico fluxo continuo,
    preservando para cada indice GLOBAL qual (arquivo, indice_local) ele
    representa. Usado no modo --combinado, quando um tronco-raiz comeca num
    arquivo e continua no seguinte sem repetir o separador.
    """
    from docx import Document
    paragrafos = []
    origem_por_idx = {}
    for path in docx_paths:
        doc = Document(str(path))
        paras = doc.paragraphs
        base = len(paragrafos)
        for local_i in range(len(paras)):
            origem_por_idx[base + local_i] = (path.name, local_i)
        paragrafos.extend(paras)
    return paragrafos, origem_por_idx


def main(argv=None):
    argv = argv or sys.argv[1:]
    aqui = Path(__file__).resolve().parent

    if argv and argv[0] == "--combinado":
        resto = argv[1:]
        if len(resto) < 2:
            print("Uso: python extrator.py --combinado <docx1> <docx2> [... <docxN>] [dir_saida]")
            sys.exit(1)
        # ultimo argumento e dir_saida so se NAO terminar em .docx
        if resto[-1].lower().endswith(".docx"):
            docx_args, out_dir = resto, aqui
        else:
            docx_args, out_dir = resto[:-1], Path(resto[-1])
        docx_paths = [Path(a) for a in docx_args]

        paragrafos, origem_por_idx = _ler_paragrafos_com_origem(docx_paths)
        arquivo_origem = " + ".join(p.name for p in docx_paths)
        ext = Extrator(
            arquivo_origem=arquivo_origem,
            origem_por_idx=origem_por_idx,
            arquivos_origem=[p.name for p in docx_paths],
        )
        ext.processar(paragrafos)
    else:
        docx_path = Path(argv[0]) if len(argv) >= 1 else aqui / "originais" / "primeira_parte_limpo.docx"
        out_dir = Path(argv[1]) if len(argv) >= 2 else aqui

        from docx import Document
        doc = Document(str(docx_path))
        paragrafos = doc.paragraphs
        # arquivo unico: indice local == indice global, mas ja no formato
        # padrao (arquivo_origem + paragrafo_idx_local) para manter o mesmo
        # esquema de dados independente do modo usado.
        origem_por_idx = {i: (docx_path.name, i) for i in range(len(paragrafos))}
        ext = Extrator(
            arquivo_origem=docx_path.name,
            origem_por_idx=origem_por_idx,
            arquivos_origem=[docx_path.name],
        )
        ext.processar(paragrafos)

    dados = ext.to_dict()
    (out_dir / "dados.json").write_text(
        json.dumps(dados, ensure_ascii=False, indent=2), encoding="utf-8")
    (out_dir / "relatorio_extracao.md").write_text(
        relatorio_extracao(ext), encoding="utf-8")
    (out_dir / "relatorio_ambiguidades.md").write_text(
        relatorio_ambiguidades(ext), encoding="utf-8")

    print(f"OK: {len(ext.troncos_raiz)} troncos-raiz, "
          f"{sum(len(r.troncos_familiares) for r in ext.troncos_raiz)} troncos familiares, "
          f"{len(ext.ambiguidades)} ambiguidades.")
    print(f"Saída em: {out_dir}")


if __name__ == "__main__":
    main()
